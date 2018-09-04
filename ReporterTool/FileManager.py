from docx import Document

class WordFileManager:

    def __init__(self, path):
        self.document = path

    def setFilePath(self, path):
        self.document = path

    def readCells(self):
        wordDoc = Document(self.document)
        rows = []
        for table in wordDoc.tables:
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
        return rows

    def getRedRows(self):
        wordDoc = Document(self.document)
        rows = []
        for table in wordDoc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if("Error:" in cell.text):
                        rows.append([cell.text for cell in row.cells])

        return rows

    def readText(self):
        wordDoc = Document(self.document)
        paras = []
        paras.append([p.text for p in wordDoc.paragraphs])
        return paras

class FileManager:

    def __init__(self, path):
        self.document = path

    def readFile(self):
        file = open(self.document, "r")
        return file.readlines()
<<<<<<< HEAD
=======

    def getErrorID(self, errorCode):
        file = open(self.document, "r")
        source = self.readFile()
        errorID = {}
        for line in source:
            if (errorCode in line.lower()):
                err = line[line.lower().find(errorCode) + len(errorCode):-3]
                err = err.strip()
                errorID.update({line:err})
        return errorID
>>>>>>> 52363c1aa4428de7fd1f40b4c9661c4721f85f4c
